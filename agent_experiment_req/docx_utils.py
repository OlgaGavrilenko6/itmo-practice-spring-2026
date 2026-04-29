# import re
# from dataclasses import dataclass
# from pathlib import Path
# from typing import Optional, Union, Iterable
# from docx import Document
# from docx.text.paragraph import Paragraph
#
#
# DEFAULT_INVALID_HEADINGS = [
#     '1 Область применения',
#     '2 Нормативные ссылки',
#     '3 Термины, определения и сокращения',
#     '3 Термины и определения',
#     'Введение',
#     'Содержание',
#     'Предисловие',
#     'Библиография',
#     'Список источников',
#     'Список литературы',
# ]
#
#
# @dataclass
# class FilterDecision:
#     keep: bool
#     reason: str
#     score: float
#
#
# def _normalize_text(text: str) -> str:
#     if not text:
#         return ""
#     text = text.lower().replace("ё", "е")
#     text = re.sub(r"\s+", " ", text)
#     return text.strip()
#
#
# def _clean_optional_text(value: Optional[str]) -> Optional[str]:
#     if value is None:
#         return None
#     cleaned = value.strip()
#     return cleaned or None
#
#
# def read_docx_paragraphs(file_path: Union[str, Path]) -> list[str]:
#     doc = Document(str(file_path))
#     paragraphs: list[str] = []
#
#     for p in doc.paragraphs:
#         text = p.text.strip()
#         if text:
#             paragraphs.append(text)
#
#     return paragraphs
#
#
# HARD_RULES = [
#     (re.compile(r'^\s*\((измененн\w*\s+редакц\w*|исключен\w*|утратил\w*\s+силу).*?\)\.?\s*$', re.I), "editorial_note_only"),
#     (re.compile(r'^\s*\d+(?:\.\d+)*\s*\(исключен\w*.*?\)\.?\s*$', re.I), "excluded_clause"),
#     (re.compile(r'^\s*(таблица|рисунок|приложение)\b.*$', re.I), "table_figure_appendix_header"),
#     (re.compile(r'^\s*(\d+(?:\.\d+)*|№\s*\d+)\s*$', re.I), "bare_number"),
# ]
#
# REQ_CUES = re.compile(
#     r'\b(должн\w*|следует|не допускает\w*|не допускается|необходимо|предусматрива\w*|обеспечива\w*|запреща\w*)\b',
#     re.I,
# )
#
#
# def soft_score(text: str) -> float:
#     t = text.strip()
#     if not t:
#         return 999.0
#
#     letters = sum(ch.isalpha() for ch in t)
#     digits = sum(ch.isdigit() for ch in t)
#     punct = sum(ch in ".,;:()[]№–-/" for ch in t)
#
#     ratio_letters = letters / max(1, len(t))
#     ratio_digits = digits / max(1, len(t))
#     ratio_punct = punct / max(1, len(t))
#
#     score = 0.0
#     if len(t) < 30:
#         score += 1.5
#     if ratio_letters < 0.45:
#         score += 2.0
#     if ratio_digits > 0.20:
#         score += 1.0
#     if ratio_punct > 0.18:
#         score += 0.8
#
#     if re.fullmatch(
#         r'[\s,;]*(СП|ГОСТ|СНиП)\s*\d+(\.\d+)*[\s,;]*((СП|ГОСТ|СНиП)\s*\d+(\.\d+)*)*[\s,;]*\.?',
#         t,
#         flags=re.I,
#     ):
#         score += 2.5
#
#     if REQ_CUES.search(t):
#         score -= 2.0
#
#     return score
#
#
# def filter_paragraph(text: str) -> FilterDecision:
#     t = text.strip()
#     for rule, name in HARD_RULES:
#         if rule.match(t):
#             return FilterDecision(keep=False, reason=name, score=999.0)
#
#     score = soft_score(t)
#     if score >= 3.0:
#         return FilterDecision(keep=False, reason="soft_garbage", score=score)
#     if score >= 2.0:
#         return FilterDecision(keep=True, reason="quarantine_candidate", score=score)
#     return FilterDecision(keep=True, reason="ok", score=score)
#
#
# def _detect_heading_level(paragraph: Paragraph) -> Optional[int]:
#     style_name = (paragraph.style.name or "").strip().lower()
#
#     match = re.search(r'(heading|заголовок)\s*(\d+)', style_name)
#     if match:
#         return int(match.group(2))
#
#     text = paragraph.text.strip()
#     numbered = re.match(r'^(\d+(?:\.\d+)*)\s+\S', text)
#     if numbered:
#         return min(numbered.group(1).count('.') + 1, 6)
#
#     return None
#
#
# def _is_invalid_heading(text: str, invalid_headings_norm: set[str]) -> bool:
#     norm = _normalize_text(text)
#     if norm in invalid_headings_norm:
#         return True
#
#     return any(norm.startswith(bad) for bad in invalid_headings_norm)
#
#
# def _format_heading(prefix_level: int, heading_text: str) -> str:
#     return f"{'#' * max(prefix_level, 1)} {heading_text.strip()}"
#
#
# def _build_fragment(current_headings: dict[int, str], body_lines: list[str]) -> Optional[str]:
#     cleaned_body = [line.strip() for line in body_lines if line and line.strip()]
#     if not cleaned_body:
#         return None
#
#     heading_lines = [
#         _format_heading(level, current_headings[level])
#         for level in sorted(current_headings)
#         if current_headings.get(level)
#     ]
#
#     parts = heading_lines + cleaned_body
#     text = "\n".join(parts).strip()
#     return text or None
#
#
# def parse_docx_to_fragments(
#     file_path: Union[str, Path],
#     invalid_headings: Optional[Iterable[str]] = None,
#     extract_tables: bool = False,
# ) -> list[str]:
#     if extract_tables:
#         raise NotImplementedError("extract_tables=True пока не поддержан в этой версии")
#
#     doc = Document(str(file_path))
#     invalid_headings = list(invalid_headings or DEFAULT_INVALID_HEADINGS)
#     invalid_headings_norm = {_normalize_text(x) for x in invalid_headings if x}
#
#     fragments: list[str] = []
#     current_headings: dict[int, str] = {}
#     body_lines: list[str] = []
#     skip_current_branch = False
#
#     def flush_current() -> None:
#         fragment = _build_fragment(current_headings, body_lines)
#         if fragment:
#             fragments.append(fragment)
#         body_lines.clear()
#
#     for paragraph in doc.paragraphs:
#         text = paragraph.text.strip()
#         if not text:
#             continue
#
#         heading_level = _detect_heading_level(paragraph)
#         if heading_level is not None:
#             flush_current()
#
#             current_headings = {level: value for level, value in current_headings.items() if level < heading_level}
#             current_headings[heading_level] = text
#             skip_current_branch = any(
#                 _is_invalid_heading(value, invalid_headings_norm)
#                 for value in current_headings.values()
#             )
#             continue
#
#         if skip_current_branch:
#             continue
#
#         body_lines.append(text)
#
#     flush_current()
#     return fragments
#
#
# def extract_requirement_fragments_from_docx(
#     file_path: Union[str, Path],
#     invalid_headings: Optional[Iterable[str]] = None,
#     extract_tables: bool = False,
# ) -> list[str]:
#     raw_fragments = parse_docx_to_fragments(
#         file_path=file_path,
#         invalid_headings=invalid_headings,
#         extract_tables=extract_tables,
#     )
#     return [fragment for fragment in raw_fragments if filter_paragraph(fragment).keep]
#
#
# def extract_named_chapter_from_docx(
#     file_path: Union[str, Path],
#     chapter_number: str,
#     chapter_title_pattern: str,
#     stop_on_next_top_level_heading: bool = True,
# ) -> Optional[str]:
#     paragraphs = read_docx_paragraphs(file_path)
#     if not paragraphs:
#         return None
#
#     start_idx: Optional[int] = None
#     title_pattern_norm = _normalize_text(chapter_title_pattern)
#
#     for i, paragraph in enumerate(paragraphs):
#         norm = _normalize_text(paragraph)
#
#         if re.match(rf"^{re.escape(chapter_number)}\b", paragraph.strip()) and title_pattern_norm in norm:
#             start_idx = i
#             break
#
#         if title_pattern_norm in norm and start_idx is None:
#             if chapter_number == "1" and "область применения" in norm:
#                 start_idx = i
#                 break
#             if chapter_number == "3" and "термины" in norm:
#                 start_idx = i
#                 break
#
#     if start_idx is None:
#         return None
#
#     collected: list[str] = []
#
#     for i in range(start_idx, len(paragraphs)):
#         p = paragraphs[i]
#
#         if i > start_idx and stop_on_next_top_level_heading:
#             if re.match(r"^\s*\d+\s+[А-ЯA-ZЁ]", p):
#                 break
#
#         collected.append(p)
#
#     text = "\n".join(collected).strip()
#     return text or None
#
#
# def build_global_contexts_from_docx(
#     file_path: Union[str, Path],
# ) -> tuple[Optional[str], Optional[str]]:
#     scope_text = extract_named_chapter_from_docx(
#         file_path=file_path,
#         chapter_number="1",
#         chapter_title_pattern="область применения",
#     )
#
#     terms_text = extract_named_chapter_from_docx(
#         file_path=file_path,
#         chapter_number="3",
#         chapter_title_pattern="термины",
#     )
#
#     return scope_text, terms_text
#
#
# def combine_global_context(
#     scope_text: Optional[str],
#     terms_text: Optional[str],
# ) -> str:
#     parts: list[str] = []
#
#     if scope_text:
#         parts.append(scope_text)
#
#     if terms_text:
#         parts.append(terms_text)
#
#     return "\n\n".join(parts).strip()

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Optional, Union, Iterable, Sequence

from docx import Document
from docx.text.paragraph import Paragraph


DEFAULT_INVALID_HEADINGS = [
    "1 Область применения",
    "2 Нормативные ссылки",
    "3 Термины, определения и сокращения",
    "3 Термины и определения",
    "Введение",
    "Содержание",
    "Предисловие",
    "Библиография",
    "Список источников",
    "Список литературы",
]


@dataclass
class FilterDecision:
    keep: bool
    reason: str
    score: float


def _normalize_text(text: str) -> str:
    if not text:
        return ""

    text = text.lower().replace("ё", "е")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _clean_optional_text(value: Optional[str]) -> Optional[str]:
    if value is None:
        return None

    cleaned = value.strip()
    return cleaned or None


def read_docx_paragraphs(file_path: Union[str, Path]) -> list[str]:
    doc = Document(str(file_path))
    paragraphs: list[str] = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append(text)

    return paragraphs


HARD_RULES = [
    (
        re.compile(
            r"^\s*\((измененн\w*\s+редакц\w*|исключен\w*|утратил\w*\s+силу).*?\)\.?\s*$",
            re.I,
        ),
        "editorial_note_only",
    ),
    (
        re.compile(
            r"^\s*\d+(?:\.\d+)*\s*\(исключен\w*.*?\)\.?\s*$",
            re.I,
        ),
        "excluded_clause",
    ),
    (
        re.compile(r"^\s*(таблица|рисунок|приложение)\b.*$", re.I),
        "table_figure_appendix_header",
    ),
    (
        re.compile(r"^\s*(\d+(?:\.\d+)*|№\s*\d+)\s*$", re.I),
        "bare_number",
    ),
]


REQ_CUES = re.compile(
    r"\b(должн\w*|следует|не допускает\w*|не допускается|необходимо|предусматрива\w*|обеспечива\w*|запреща\w*)\b",
    re.I,
)


def soft_score(text: str) -> float:
    t = text.strip()
    if not t:
        return 999.0

    letters = sum(ch.isalpha() for ch in t)
    digits = sum(ch.isdigit() for ch in t)
    punct = sum(ch in ".,;:()[]№–-/" for ch in t)

    ratio_letters = letters / max(1, len(t))
    ratio_digits = digits / max(1, len(t))
    ratio_punct = punct / max(1, len(t))

    score = 0.0

    if len(t) < 30:
        score += 1.5

    if ratio_letters < 0.45:
        score += 2.0

    if ratio_digits > 0.20:
        score += 1.0

    if ratio_punct > 0.18:
        score += 0.8

    if re.fullmatch(
        r"[\s,;]*(СП|ГОСТ|СНиП)\s*\d+(\.\d+)*[\s,;]*((СП|ГОСТ|СНиП)\s*\d+(\.\d+)*)*[\s,;]*\.?",
        t,
        flags=re.I,
    ):
        score += 2.5

    if REQ_CUES.search(t):
        score -= 2.0

    return score


def filter_paragraph(text: str) -> FilterDecision:
    t = text.strip()

    for rule, name in HARD_RULES:
        if rule.match(t):
            return FilterDecision(
                keep=False,
                reason=name,
                score=999.0,
            )

    score = soft_score(t)

    if score >= 3.0:
        return FilterDecision(
            keep=False,
            reason="soft_garbage",
            score=score,
        )

    if score >= 2.0:
        return FilterDecision(
            keep=True,
            reason="quarantine_candidate",
            score=score,
        )

    return FilterDecision(
        keep=True,
        reason="ok",
        score=score,
    )


def _detect_heading_level(paragraph: Paragraph) -> Optional[int]:
    style_name = (paragraph.style.name or "").strip().lower()

    match = re.search(r"(heading|заголовок)\s*(\d+)", style_name)
    if match:
        return int(match.group(2))

    text = paragraph.text.strip()
    numbered = re.match(r"^(\d+(?:\.\d+)*)\s+\S", text)

    if numbered:
        return min(numbered.group(1).count(".") + 1, 6)

    return None


def _is_invalid_heading(text: str, invalid_headings_norm: set[str]) -> bool:
    norm = _normalize_text(text)

    if norm in invalid_headings_norm:
        return True

    return any(norm.startswith(bad) for bad in invalid_headings_norm)


def _format_heading(prefix_level: int, heading_text: str) -> str:
    return f"{'#' * max(prefix_level, 1)} {heading_text.strip()}"


def _build_fragment(
    current_headings: dict[int, str],
    body_lines: list[str],
) -> Optional[str]:
    cleaned_body = [
        line.strip()
        for line in body_lines
        if line and line.strip()
    ]

    if not cleaned_body:
        return None

    heading_lines = [
        _format_heading(level, current_headings[level])
        for level in sorted(current_headings)
        if current_headings.get(level)
    ]

    parts = heading_lines + cleaned_body
    text = "\n".join(parts).strip()

    return text or None


def parse_docx_to_fragments(
    file_path: Union[str, Path],
    invalid_headings: Optional[Iterable[str]] = None,
    extract_tables: bool = False,
) -> list[str]:
    if extract_tables:
        raise NotImplementedError("extract_tables=True пока не поддержан в этой версии")

    doc = Document(str(file_path))

    invalid_headings = list(invalid_headings or DEFAULT_INVALID_HEADINGS)
    invalid_headings_norm = {
        _normalize_text(x)
        for x in invalid_headings
        if x
    }

    fragments: list[str] = []
    current_headings: dict[int, str] = {}
    body_lines: list[str] = []
    skip_current_branch = False

    def flush_current() -> None:
        fragment = _build_fragment(current_headings, body_lines)
        if fragment:
            fragments.append(fragment)
        body_lines.clear()

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if not text:
            continue

        heading_level = _detect_heading_level(paragraph)

        if heading_level is not None:
            flush_current()

            current_headings = {
                level: value
                for level, value in current_headings.items()
                if level < heading_level
            }

            current_headings[heading_level] = text

            skip_current_branch = any(
                _is_invalid_heading(value, invalid_headings_norm)
                for value in current_headings.values()
            )

            continue

        if skip_current_branch:
            continue

        body_lines.append(text)

    flush_current()

    return fragments


CLAUSE_ID_RE = re.compile(
    r"(?m)^\s*#*\s*(\d+(?:\.\d+)+)\b"
)


def extract_clause_ids_from_fragment(fragment: str) -> set[str]:
    if not fragment:
        return set()

    return {
        match.group(1)
        for match in CLAUSE_ID_RE.finditer(fragment)
    }


def filter_fragments_by_clause_ids(
    fragments: Sequence[str],
    selected_clause_ids: Sequence[str],
) -> list[str]:
    selected = set(selected_clause_ids)

    if not selected:
        return list(fragments)

    filtered: list[str] = []

    for fragment in fragments:
        fragment_clause_ids = extract_clause_ids_from_fragment(fragment)

        if fragment_clause_ids & selected:
            filtered.append(fragment)

    return filtered


def extract_requirement_fragments_from_docx(
    file_path: Union[str, Path],
    invalid_headings: Optional[Iterable[str]] = None,
    extract_tables: bool = False,
    selected_clause_ids: Optional[Sequence[str]] = None,
) -> list[str]:
    raw_fragments = parse_docx_to_fragments(
        file_path=file_path,
        invalid_headings=invalid_headings,
        extract_tables=extract_tables,
    )

    kept_fragments = [
        fragment
        for fragment in raw_fragments
        if filter_paragraph(fragment).keep
    ]

    if selected_clause_ids:
        kept_fragments = filter_fragments_by_clause_ids(
            fragments=kept_fragments,
            selected_clause_ids=selected_clause_ids,
        )

    return kept_fragments


def extract_named_chapter_from_docx(
    file_path: Union[str, Path],
    chapter_number: str,
    chapter_title_pattern: str,
    stop_on_next_top_level_heading: bool = True,
) -> Optional[str]:
    paragraphs = read_docx_paragraphs(file_path)

    if not paragraphs:
        return None

    start_idx: Optional[int] = None
    title_pattern_norm = _normalize_text(chapter_title_pattern)

    for i, paragraph in enumerate(paragraphs):
        norm = _normalize_text(paragraph)

        if (
            re.match(rf"^{re.escape(chapter_number)}\b", paragraph.strip())
            and title_pattern_norm in norm
        ):
            start_idx = i
            break

        if title_pattern_norm in norm and start_idx is None:
            if chapter_number == "1" and "область применения" in norm:
                start_idx = i
                break

            if chapter_number == "3" and "термины" in norm:
                start_idx = i
                break

    if start_idx is None:
        return None

    collected: list[str] = []

    for i in range(start_idx, len(paragraphs)):
        p = paragraphs[i]

        if i > start_idx and stop_on_next_top_level_heading:
            if re.match(r"^\s*\d+\s+[А-ЯA-ZЁ]", p):
                break

        collected.append(p)

    text = "\n".join(collected).strip()

    return text or None


def build_global_contexts_from_docx(
    file_path: Union[str, Path],
) -> tuple[Optional[str], Optional[str]]:
    scope_text = extract_named_chapter_from_docx(
        file_path=file_path,
        chapter_number="1",
        chapter_title_pattern="область применения",
    )

    terms_text = extract_named_chapter_from_docx(
        file_path=file_path,
        chapter_number="3",
        chapter_title_pattern="термины",
    )

    return scope_text, terms_text


def combine_global_context(
    scope_text: Optional[str],
    terms_text: Optional[str],
) -> str:
    parts: list[str] = []

    if scope_text:
        parts.append(scope_text)

    if terms_text:
        parts.append(terms_text)

    return "\n\n".join(parts).strip()